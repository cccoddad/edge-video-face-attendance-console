from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REFERENCE = Path(r"C:\Users\CC\Desktop\春秋招聘\简历项目\RK3568实时音视频边缘分析网关_低预算完整实施手册.docx")
OUTPUT = ROOT / "docs" / "基于RTSP边缘视频源的离线人脸识别与考勤控制台_完整实施方案.docx"


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_para(doc, text="", *, bold_prefix=None, color=None, before=0, after=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.right_indent = Cm(0.45)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        p_pr.append(shd)
        r = p.add_run(line if line else " ")
        set_run_font(r, name="Consolas", size=9)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "D9EAF7")
        if widths:
            set_cell_width(cell, widths[index])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = ""
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[index], widths[index])
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            set_run_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=15, bold=True, color=(31, 78, 121))
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=(47, 84, 150))
    return p


def clear_document_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.85)
    section.left_margin = Cm(2.08)
    section.right_margin = Cm(2.08)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name in ("Heading 1", "Heading 2", "List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.add_run("离线人脸识别与考勤控制台 · 完整实施方案"), size=8.5, color=(100, 100, 100))
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    set_run_font(footer_p.add_run("面向 C++ / Qt / 音视频求职展示 | 版本 1.0 | 2026-08"), size=8.5, color=(110, 110, 110))


def build():
    if not REFERENCE.exists():
        raise FileNotFoundError(f"Reference document not found: {REFERENCE}")
    doc = Document(REFERENCE)
    clear_document_body(doc)
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("基于 RTSP 边缘视频源的离线人脸识别与考勤控制台")
    set_run_font(r, size=22, bold=True, color=(31, 78, 121))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("与 RK3568 实时音视频边缘分析网关联动的完整升级与实施方案")
    set_run_font(r, size=13, color=(89, 89, 89))
    add_para(doc, "适用读者：具备 C/C++、Qt、OpenCV 基础，计划投递嵌入式软件、音视频、边缘 AI 或 Linux 应用开发岗位的学习者。", bold_prefix="适用读者：", before=4)
    add_para(doc, "重要说明：本方案将“已完成基线”“待实现功能”和“远期规划”严格分开。简历、答辩和演示中只能把有代码、日志、测试或视频证据的功能写成已实现成果；未完成内容只能写为设计或下一阶段计划。", bold_prefix="重要说明：", color=(156, 87, 0), after=12)

    h1(doc, "目录与阅读方式")
    for item in [
        "1. 项目最终定位与方向评估", "2. 为什么不能做成第二个音视频网关", "3. 系统边界、联动关系与项目命名",
        "4. 产品目标、非目标与核心场景", "5. 总体架构与模块职责", "6. 视频源、RTSP 接入与断流恢复",
        "7. 人脸注册、识别决策与考勤状态机", "8. SQLite 数据模型、事务与隐私边界", "9. 线程、队列、指标与错误处理",
        "10. 分阶段实施路线与验收标准", "11. 测试、演示、性能证据与仓库规范", "12. 简历表达、面试追问与最终价值判断",
        "13. 风险、成本控制与下一步执行清单"
    ]:
        add_bullet(doc, item)

    h1(doc, "1. 项目最终定位与方向评估")
    add_para(doc, "本项目不是把旧的 Qt 人脸识别程序简单换一个界面、换一个 Qt 版本，也不是在 RK3568 上复制一套采集、编码和推流逻辑。它最终定位为“身份识别与业务控制面”：从边缘视频源中稳定获得画面，将人脸特征匹配转化为可追溯、可去重、可查询的考勤事件。")
    add_para(doc, "与之对应，RK3568 实时音视频边缘分析网关承担“媒体数据面”：它处理摄像头、麦克风、V4L2、ALSA、视频时间戳、编解码、RTSP、端侧推理和故障恢复。两个项目通过 RTSP 视频流和可选状态接口联动，但各自只负责自己所在层次的问题。")
    add_table(doc, ["项目", "最终定位", "主要证明能力", "不应重复承担的内容"], [
        ["RK3568 实时音视频边缘分析网关", "边缘媒体数据面", "V4L2/ALSA、PTS、有界队列、RKNN/RGA、MPP/FFmpeg、MP4/RTSP、systemd 与恢复", "考勤业务规则、人员资料、报表与桌面管理"],
        ["离线人脸识别与考勤控制台", "身份与业务控制面", "RTSP 消费、Qt/C++ 多线程、识别决策、SQLite 事务、事件幂等、报表与审计", "V4L2 采集、AAC/H.264 编码、RTSP 服务端、MPP 推流"],
    ], [3.2, 3.2, 5.1, 5.2])
    add_para(doc, "方向判断：采用该分层后，第二项目的潜在含金量明显高于原始“Qt 人脸考勤系统”。它既能和音视频主项目形成真实的接口关系，又能展示不同的工程能力。前提是主项目和第二项目分别有独立仓库、README、测试证据和演示重点，不能把同一段工作拆成两段重复写入简历。")

    h2(doc, "1.1 含金量的真实分级")
    add_table(doc, ["形态", "估计价值", "原因"], [
        ["原始 Qt + OpenCV + 人脸 SDK + SQLite", "中低", "功能完整但接近课程项目，第三方 SDK 调用占比高，工程可靠性证据不足。"],
        ["重复迁移到 RK3568 并再做采集/推流", "中等", "技术栈更多，但与音视频网关高度同质化，容易被视为一个项目拆成两个。"],
        ["网关 + RTSP 考勤控制台的分层联动", "中高至较高", "可同时证明媒体链路、边缘视频接入、身份决策、数据一致性和桌面端交付能力。"],
    ], [3.5, 2.4, 10.8])
    add_para(doc, "任何分级都不是由项目标题决定，而是由证据决定。能够提供构建脚本、运行日志、故障恢复录像、指标表、数据库样例和设计说明，项目的可信度会显著提高；只展示一张识别成功截图，哪怕标题很长，含金量也不会高。")

    h1(doc, "2. 为什么不能做成第二个音视频网关")
    add_para(doc, "RK3568 网关手册已经覆盖了视频采集、音频采集、时间戳、队列、RKNN/RGA、MPP/FFmpeg、MP4、RTSP、监控和 systemd 恢复。若人脸考勤项目再次实现同一套内容，两个项目会在设备层、媒体层、AI 层和运维层同时重叠。")
    add_table(doc, ["重叠层次", "网关项目已承担", "第二项目正确做法"], [
        ["设备层", "V4L2、UVC 摄像头、ALSA、USB 设备枚举", "仅把本地摄像头视为开发输入；生产联动消费 RTSP。"],
        ["媒体层", "H.264/AAC、MPP、FFmpeg、MP4 与 RTSP 推流", "只负责拉取、解码和显示 RTSP 视频，不建立推流服务端。"],
        ["AI 层", "可运行 RKNN 目标检测与叠加", "聚焦人脸身份匹配的决策逻辑和结果审计，不堆叠第二套通用视觉管线。"],
        ["可靠性层", "systemd、设备恢复、网络重连、长稳监控", "实现客户端视频源状态和重连，不复制板端守护与硬件恢复。"],
    ], [3.0, 6.6, 7.1])
    add_para(doc, "因此，第二项目的关键问题应当从“如何把视频编码推到网络上”切换为“如何从可能断开的网络视频中得到可靠的身份事件”。前者是音视频工程问题，后者是视频接入、C++ 应用架构、数据一致性和业务决策问题。")

    h1(doc, "3. 系统边界、联动关系与项目命名")
    add_para(doc, "建议保持两个独立项目。网关可以在 RK3568 上运行，考勤控制台运行在 Windows PC 上；未来控制台也可移植到 Linux，但这不是第一阶段目标。二者不共享数据库，不相互依赖私有源码，只通过公开、简单的接口协作。")
    add_code(doc, "USB 摄像头 / 麦克风\n        │\n        ▼\nRK3568 音视频边缘网关 ── RTSP(H.264 视频流) ──► 人脸识别与考勤控制台\n        │                                              │\n        └── /status（可选状态接口） ────────────────────┤\n                                                       ▼\n                                  人脸匹配 → 决策状态机 → SQLite 事件账本 / CSV")
    add_para(doc, "项目名称应使用“基于 RTSP 边缘视频源的离线人脸识别与考勤控制台”。不要以“Qt 人脸识别系统”为正式名称，因为 Qt 只是界面和应用框架；也不要在名称中堆砌 RK3568、FFmpeg、YOLO、OpenCV 等不属于该项目核心职责的技术词。")
    add_para(doc, "最小联动接口只有一个 RTSP 地址，例如 rtsp://<rk3568-ip>:8554/live。网关端负责保证地址可播放；控制台负责连接、读取、解码、预览、识别和事件落库。第二阶段可增加只读状态接口，向控制台暴露网关在线状态、采集帧率、编码帧率、最近错误和当前流地址。")

    h1(doc, "4. 产品目标、非目标与核心场景")
    h2(doc, "4.1 第一版目标")
    for item in [
        "支持本地 USB 摄像头、RTSP 视频流和本地视频文件三种输入方式。",
        "支持人员注册、人员启停用、人脸模板保存和现场抓拍留证。",
        "支持人脸检测、特征匹配、阈值判断、连续帧确认和未知人员提示。",
        "支持签到、签退、重复事件冷却、当日状态查询和 CSV 导出。",
        "支持视频源断流、自动重连、错误日志、运行指标和数据库事务。",
        "支持使用 RK3568 网关 RTSP 流完成端到端演示。",
    ]:
        add_bullet(doc, item)
    h2(doc, "4.2 明确非目标")
    for item in [
        "第一版不实现可靠活体检测，不应宣传为高安全门禁或防照片攻击系统。",
        "第一版不实现复杂排班、节假日、审批流、云端 SaaS、多组织权限或多机集群。",
        "第一版不在控制台中重复实现 H.264/AAC 编码、MP4 封装、RTSP 推流与 MPP。",
        "第一版不保证多人密集场景下的精确全员考勤；检测到多人时优先提示单人站位。",
        "第一版不提交真实人员照片、人脸模板或考勤数据库到公开仓库。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "这些非目标并不是能力不足，而是有意识的范围控制。对求职项目而言，一个边界清楚、可稳定演示的单人离线考勤控制台，远胜于一个写着多路摄像头、云端同步、活体检测，却无法证明任何一项的“大而全系统”。")

    h1(doc, "5. 总体架构与模块职责")
    add_para(doc, "工程应从当前 Qt Widgets 基线逐步模块化。UI、视频输入、识别、决策、数据库和指标不能继续混在窗口类中。模块化的目的不是形式化重构，而是让每一部分都能独立测试、替换和解释。")
    add_code(doc, "app/        程序入口、依赖装配、全局配置\nmedia/      LocalCameraSource / RtspSource / VideoFileSource / SourceManager\nvision/     检测、特征提取、匹配、跟踪、识别流水线\ndomain/     人员、考勤事件、连续确认、状态机、业务规则\nstorage/    SQLite、迁移、Repository、事务、快照索引\nmonitor/    指标、健康状态、事件日志\nui/         预览、注册、查询、视频源配置、指标面板\ntests/      规则引擎、去重、时间边界、文件源回归测试")
    add_table(doc, ["模块", "输入", "输出", "职责边界"], [
        ["media", "设备号、RTSP URL、视频文件", "VideoFrame、连接状态", "只负责得到最新视频帧和管理连接，不做业务判断。"],
        ["vision", "VideoFrame、人脸模板", "识别候选结果", "只回答“谁最像、分数多少”，不直接写数据库。"],
        ["domain", "识别候选、人员状态、当前日期", "考勤事件或拒绝原因", "负责阈值、连续确认、冷却、签到签退状态机。"],
        ["storage", "人员、模板、事件、快照", "事务结果", "负责持久化、迁移、唯一性和审计。"],
        ["monitor", "帧率、耗时、错误、重连", "指标和日志", "让运行状态可见、可诊断。"],
        ["ui", "用户操作、状态信号", "配置命令、界面渲染", "不执行阻塞 I/O 或重计算。"],
    ], [2.3, 4.0, 3.8, 6.6])

    h1(doc, "6. 视频源、RTSP 接入与断流恢复")
    add_para(doc, "视频源是两个项目联动的最小切入点。应定义统一的 IVideoSource 接口，使本地摄像头、RTSP 和视频文件都可产生同一种 VideoFrame。这样既能在没有硬件时用视频文件做测试，也能在摄像头插在 RK3568 上时直接消费板端的 RTSP 流。")
    add_code(doc, "struct VideoFrame {\n    uint64_t sequence;\n    int64_t receive_time_ms;\n    QString source_id;\n    cv::Mat image;\n};\n\nclass IVideoSource {\npublic:\n    virtual bool start() = 0;\n    virtual void stop() = 0;\n    virtual SourceState state() const = 0;\n};")
    h2(doc, "6.1 视频源状态机")
    add_code(doc, "IDLE → CONNECTING → STREAMING\n                    │\n                    ├── 读取超时 / 解码失败 → RECONNECT_WAIT → CONNECTING\n                    │\n                    └── 重试超过上限 → DEGRADED（等待用户操作或后台探测）")
    add_para(doc, "界面不能把 RTSP 断流显示成“识别失败”。识别不到人说明视频仍在输入但画面中没有满足条件的人脸；视频源断流说明输入链路不可用。这两类情况应有不同颜色、不同日志和不同业务处理。断流期间不产生考勤事件，也不清空历史考勤状态。")
    add_table(doc, ["场景", "界面表现", "后台行为", "考勤行为"], [
        ["RTSP 地址错误", "连接失败与错误原因", "停止读取，允许修改地址", "不产生事件"],
        ["网关短暂重启", "显示重连中与次数", "按间隔重试", "不产生事件"],
        ["恢复收到视频帧", "显示已连接与 FPS", "重置错误计数", "恢复识别"],
        ["本地摄像头被占用", "显示设备不可用", "不阻塞 UI", "不产生事件"],
        ["视频文件结束", "显示回放完成", "停止文件源或允许循环", "按配置停止或演示重放"],
    ], [3.1, 4.8, 4.8, 4.0])

    h1(doc, "7. 人脸注册、识别决策与考勤状态机")
    h2(doc, "7.1 注册策略")
    add_para(doc, "人员注册不应只把一张照片和一个工号写进数据库。建议在单人、清晰、尺寸合格的前提下采集多张有效样本，例如正脸与轻微左右转脸。第一版可以保存多份模板，不必立刻实现复杂的模板融合；后续可根据实测误识率决定是否引入质量评分和最佳模板选择。")
    add_para(doc, "注册前应检查工号唯一性、人脸是否存在、画面中是否只有一人、脸部尺寸是否足够、图像是否严重模糊。成功注册时应在同一事务中保存人员资料、模板索引和照片路径；真实图片本体保存在应用数据目录，不进入仓库。")
    h2(doc, "7.2 识别结果与业务决策解耦")
    add_code(doc, "视频帧\n  → 人脸检测\n  → 特征提取\n  → 最近模板匹配（得到候选人员 + 相似度）\n  → 连续帧确认\n  → 考勤规则与状态机\n  → 事务写入考勤事件")
    add_table(doc, ["识别情况", "决策", "是否落库"], [
        ["未检测到人脸", "保持预览，清理过期候选", "否"],
        ["相似度低于阈值", "显示未知人员并记录可选运行计数", "不写正式考勤"],
        ["仅单帧匹配成功", "进入待确认状态", "否"],
        ["连续 N 帧匹配同一人员", "进入考勤规则判断", "视状态决定"],
        ["多人同时出现", "提示单人站位，暂停自动考勤", "否"],
        ["人员已禁用", "提示人员不可用", "不写正式考勤"],
        ["同类事件处于冷却时间", "抑制重复，记录统计", "可选写系统日志"],
    ], [4.1, 7.8, 4.8])
    h2(doc, "7.3 考勤状态机")
    add_code(doc, "未签到 ── 首次稳定确认 ──► 已签到\n已签到 ── 达到签退窗口 ──► 已签退\n已签到 ── 未到签退窗口 ──► 忽略重复识别\n已签退 ── 再次出现 ──► 默认不生成常规事件")
    add_para(doc, "第一版建议先做单班次场景：当天首次稳定确认生成签到；进入配置的签退窗口后稳定确认生成签退；其余识别只更新界面或记录为被抑制事件。跨天班、审批、补签、节假日和复杂排班应被列为后续功能，而不是占用第一版主线。")

    h1(doc, "8. SQLite 数据模型、事务与隐私边界")
    add_para(doc, "SQLite 适合该项目的离线单机特性：不需要安装服务端，便于发布和备份，能够支持人员、模板、考勤事件和日志。为了从“简单插入记录”升级为可靠的事件账本，建议拆分以下数据表。")
    add_table(doc, ["数据表", "关键字段", "用途"], [
        ["person", "employee_no、name、department、status", "人员基础资料；employee_no 建唯一索引。"],
        ["face_template", "person_id、feature_blob、quality_score、photo_path", "一个人可保存多个模板与样本索引。"],
        ["attendance_event", "event_uuid、person_id、event_type、event_time、score、snapshot_path", "不可随意覆盖的正式考勤事件账本。"],
        ["attendance_daily_state", "person_id、business_date、check_in_event_id、check_out_event_id", "快速判断当天签到与签退状态。"],
        ["video_source", "name、type、url_or_device、enabled", "本地设备、RTSP 与文件源配置。"],
        ["system_event", "event_time、level、module、message、error_code", "断流、重连、模型错误、数据库异常等运行记录。"],
        ["schema_version", "version、applied_at", "数据库迁移版本，避免升级破坏旧数据。"],
    ], [3.1, 7.4, 6.2])
    add_para(doc, "正式考勤写入至少应包含“插入 attendance_event、更新 attendance_daily_state、保存抓拍路径索引、写入必要系统日志”四类动作。它们应在同一个 SQLite 事务中提交。对 event_uuid 或业务唯一键建立唯一约束，保证程序崩溃重试时不会重复插入同一条考勤。所有 SQL 使用预编译参数，不拼接用户输入。")
    h2(doc, "8.1 隐私与数据安全边界")
    for item in [
        "测试仓库使用虚构人员、授权测试人员或占位数据；不提交真实人脸模板和完整考勤库。",
        "抓拍图片和 SQLite 数据库存放在应用数据目录，路径通过配置统一管理。",
        "README 说明数据保留策略，例如仅保留近期快照并支持管理员清理。",
        "项目说明中明确当前不具备可靠活体检测，不将其宣传为高安全门禁。",
        "演示视频中对非测试人员做人脸打码，避免无授权泄露。",
    ]:
        add_bullet(doc, item)

    h1(doc, "9. 线程、队列、指标与错误处理")
    add_para(doc, "Qt UI 线程只负责响应用户操作和刷新界面；打开视频、读取 RTSP、执行人脸特征提取、数据库事务和文件写入都不能长期占用 UI 线程。当前工程已经建立了基础工作线程边界，下一阶段应把职责进一步拆清。")
    add_table(doc, ["执行单元", "主要职责", "关键约束"], [
        ["UI 线程", "配置、预览刷新、人员管理、查询显示", "不得执行阻塞读取、模型推理或数据库大查询。"],
        ["VideoSourceWorker", "读取本地设备、RTSP 或文件，输出最新帧", "使用超时；断流后可停止并重连；队列不无限增长。"],
        ["RecognitionWorker", "检测、特征提取、模板比对", "可降低采样频率；结果带 sequence 和 source_id。"],
        ["DecisionWorker", "连续确认、阈值、冷却、状态机", "纯业务逻辑，尽量可单元测试。"],
        ["StorageWorker", "SQLite 事务、快照与日志", "失败可见；不在 UI 线程写盘。"],
        ["MetricsWorker", "帧率、耗时、错误、重连统计", "定时汇总并向 UI 发送轻量状态。"],
    ], [3.3, 7.2, 6.2])
    add_para(doc, "视频分析应采用“最新帧优先”的队列策略。若识别慢于采集，继续堆积旧帧会导致用户看到数秒前的画面，并在恢复时写入过期事件。第一版可把采集到识别的队列控制在 1～2 帧，队列满时丢弃旧帧，只让识别线程处理最新画面。")
    add_para(doc, "建议记录 Preview FPS、Recognition FPS、平均识别耗时、P95 识别耗时、RTSP 重连次数、视频源离线时长、未知人员次数、重复事件抑制次数、数据库写入失败次数。不要在未实测前承诺具体毫秒数或高帧率；先实现可测量，再用数据决定优化方向。")

    h1(doc, "10. 分阶段实施路线与验收标准")
    add_table(doc, ["阶段", "目标", "主要工作", "可验证验收"], [
        ["阶段 0：基线", "让旧项目可维护、可运行", "路径配置、模型检查、数据库修复、线程基础修复、构建与发布脚本", "可编译、可启动、可初始化 SQLite、独立发布包可运行。"],
        ["阶段 1：视频源", "接入 RTSP 并解耦输入", "IVideoSource、本地/RTSP/文件三种适配、状态页、重连", "可切换输入；断流不假死；恢复后继续预览。"],
        ["阶段 2：识别决策", "降低误识与重复事件", "阈值、连续确认、多人拦截、冷却、未知人员状态", "单帧不直接考勤；人员停留不重复打卡。"],
        ["阶段 3：业务账本", "形成可靠考勤闭环", "迁移、事务、状态机、抓拍、查询和 CSV", "重启后状态正确；重复写入被拦截；可导出。"],
        ["阶段 4：网关联动", "完成端到端演示", "对接 RK3568 RTSP、可选状态接口、故障联动", "网关恢复后控制台恢复；断流期间不误记。"],
        ["阶段 5：工程证据", "让项目可投递、可答辩", "测试、日志、指标、README、视频、简历素材", "新环境可复现；有完整和故障演示。"],
    ], [2.4, 2.8, 6.0, 5.5])
    add_para(doc, "推进原则是一次只完成一个可验收闭环。当前最适合开始的是阶段 1：视频源抽象与 RTSP 接入。它既不依赖先完成所有考勤业务，也能立刻解决“摄像头插在 RK3568 上、Windows 项目看不到”的现实问题。")

    h1(doc, "11. 测试、演示、性能证据与仓库规范")
    h2(doc, "11.1 必测场景")
    add_table(doc, ["类别", "最小测试集"], [
        ["视频输入", "本地摄像头、RTSP 正常连接、错误地址、服务端停止、网络恢复、视频文件回放完成。"],
        ["人员与识别", "重复工号、已知人员、未知人员、低置信度、多人画面、人员禁用。"],
        ["考勤逻辑", "首次签到、签退窗口、同人连续停留、跨应用重启、同日重复事件。"],
        ["存储", "数据库不可写、快照写入失败、迁移升级、CSV 导出、真实数据不入仓库。"],
        ["联动故障", "网关摄像头断开、RTSP 服务停止、网线断开、网关恢复、控制台重连。"],
    ], [2.7, 14.0])
    h2(doc, "11.2 演示证据清单")
    for item in [
        "一张双项目联动架构图：摄像头 → RK3568 网关 → RTSP → 考勤控制台 → SQLite 事件账本。",
        "一段正常演示视频：选择 RTSP、预览、注册、稳定识别、签到、查询、CSV 导出。",
        "一段故障演示视频：停止 RTSP 或拔出摄像头后显示断流，恢复后自动重连。",
        "一张指标表：预览帧率、识别帧率、平均/P95 识别耗时、重连次数、重复事件抑制次数。",
        "一份 README：依赖、模型放置方式、构建、运行、数据目录、配置、已知限制、复现步骤。",
        "一组自动化或半自动回归测试：至少能用视频文件测试规则引擎和重复事件抑制。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "仓库中应忽略运行数据库、照片、抓拍、构建目录、模型文件和本地依赖路径；保留 third_party.pri.example、配置样例和脚本。模型或 SDK 的许可证需要单独阅读，不能默认把第三方二进制和模型全部上传到公开仓库。")

    h1(doc, "12. 简历表达、面试追问与最终价值判断")
    h2(doc, "12.1 简历表达模板（仅在相应功能完成后使用）")
    add_para(doc, "项目：基于 RTSP 边缘视频源的离线人脸识别与考勤控制台", bold_prefix="项目：", before=4)
    for item in [
        "基于 C++/Qt 构建离线考勤控制台，抽象本地摄像头、RTSP 流和视频文件三类输入，支持视频源状态展示、断流检测和自动重连。",
        "将人脸检测、特征匹配与考勤决策解耦，实现相似度阈值、连续帧确认、多人拦截和重复事件冷却，避免单帧误识和人员停留造成的重复打卡。",
        "基于 SQLite 设计人员、模板、考勤事件和当日状态模型，使用事务与唯一事件标识保证事件幂等，支持抓拍留证、条件查询和 CSV 导出。",
        "接入 RK3568 音视频网关输出的 RTSP 流完成端到端演示，将边缘视频输入转化为本地可审计的身份与考勤事件。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "如果 RTSP、自动重连、连续帧确认或事务账本尚未实现，必须删除对应简历条目。项目介绍中最忌讳把设计文档、预期指标和未完成模块写成既成事实。")
    h2(doc, "12.2 必须能够回答的追问")
    add_table(doc, ["问题", "应回答的核心"], [
        ["为什么控制台不自己推 RTSP？", "网关负责媒体生产，控制台负责消费和业务决策；分层避免重复并降低耦合。"],
        ["为什么单帧匹配不能直接考勤？", "单帧受模糊、遮挡和相似度波动影响大，连续确认能降低偶发误判。"],
        ["如何避免重复打卡？", "业务状态机加冷却时间，数据库唯一事件标识提供最终兜底。"],
        ["RTSP 断流和未识别人脸有什么区别？", "前者是输入源异常，后者是输入正常但没有合格匹配；状态和业务处理不同。"],
        ["为什么用 SQLite？", "单机离线、部署简单、事务足够，符合本项目规模；不为堆技术强上服务端数据库。"],
        ["人脸安全能力有什么边界？", "当前没有可靠活体检测，面向低风险离线考勤演示，不宣传为高安全门禁。"],
    ], [4.5, 12.2])
    add_para(doc, "最终价值判断：原始项目只能说明你做过一个 Qt 桌面人脸识别应用；升级后的项目则能说明你理解边缘视频接入、线程边界、识别可信度、离线数据一致性和业务闭环。再与 RK3568 音视频网关组合后，可以形成“底层媒体数据面 + 上层身份业务控制面”的完整能力叙事。")

    h1(doc, "13. 风险、成本控制与下一步执行清单")
    h2(doc, "13.1 主要风险与对应策略")
    add_table(doc, ["风险", "表现", "策略"], [
        ["RTSP 解码兼容性", "不同流地址、编码参数或网络环境导致无法打开", "先用 ffplay/VLC 验证流，再接入控制台；保留文件源回归路径。"],
        ["人脸误识或漏识", "弱光、角度、模糊导致相似度波动", "阈值配置、连续确认、单人站位、质量过滤，记录真实样本而不凭感觉调参。"],
        ["项目范围失控", "过早加入云端、多摄像头、活体检测、复杂排班", "严格以阶段验收为准，第一版只交付单点离线闭环。"],
        ["简历可信度不足", "描述的组件多但没有运行证据", "每完成一阶段就保存代码、截图、日志、测试数据和演示片段。"],
        ["隐私数据泄露", "真实脸图、特征库被上传或录屏泄露", "使用测试数据，配置 .gitignore，演示素材打码，说明保留策略。"],
    ], [3.2, 5.6, 7.9])
    h2(doc, "13.2 下一步执行顺序")
    for item in [
        "保留当前 Qt 5.12 基线，不因升级方案而急于迁移 Qt6；先完成业务和架构升级。",
        "实现 IVideoSource、LocalCameraSource、VideoFileSource，再实现 RtspSource。",
        "增加视频源配置页、连接状态、断流提示和自动重连日志。",
        "在没有 RK3568 或摄像头时，用本地视频文件完成首轮回归测试。",
        "网关具备 RTSP 输出后，将控制台接入板端视频流，完成最小联动。",
        "随后再实现连续帧确认、签到签退状态机、SQLite 事务账本与 CSV 导出。",
        "每个阶段结束时补充 README、截图、日志和演示，不把证据工作拖到最后。",
    ]:
        add_number(doc, item)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    text = "".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "".join(cell.text for cell in row.cells)
    print(f"Created: {OUTPUT}")
    print(f"Approximate Chinese/text character count: {len(text)}")


if __name__ == "__main__":
    build()
